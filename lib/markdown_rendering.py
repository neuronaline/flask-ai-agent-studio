from __future__ import annotations

import re
from html import escape as html_escape

from docx import Document
from reportlab.lib import colors
from reportlab.lib.styles import ParagraphStyle
from reportlab.platypus import ListFlowable, ListItem, Paragraph, Preformatted, Spacer, Table, TableStyle

_MARKDOWN_FENCE_RE = re.compile(r"^\s*```([A-Za-z0-9_-]+)?\s*$")
_MARKDOWN_HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
_MARKDOWN_BULLET_RE = re.compile(r"^\s*[-*+]\s+(.*)$")
_MARKDOWN_ORDERED_RE = re.compile(r"^\s*(\d+)[.)]\s+(.*)$")
_MARKDOWN_QUOTE_RE = re.compile(r"^\s*>\s?(.*)$")
_MARKDOWN_TABLE_ROW_RE = re.compile(r"^\s*\|.*\|\s*$")
_MARKDOWN_TABLE_SEPARATOR_RE = re.compile(r"^\s*\|?(?:\s*:?-{3,}:?\s*\|)+\s*:?-{3,}:?\s*\|?\s*$")
_MARKDOWN_HORIZONTAL_RULE_RE = re.compile(r"^\s*(?:[-*_]\s*){3,}$")


def _normalize_line_endings(text: str) -> str:
    return str(text or "").replace("\r\n", "\n").replace("\r", "\n")


def _escape_pdf_text(value: str) -> str:
    return html_escape(str(value or ""), quote=False)


def _find_math_delimiter(text: str, start_index: int, delimiter: str) -> int:
    search_index = start_index
    while search_index < len(text):
        delimiter_index = text.find(delimiter, search_index)
        if delimiter_index < 0:
            return -1
        if delimiter_index > start_index and text[delimiter_index - 1] == "\\":
            search_index = delimiter_index + len(delimiter)
            continue
        if delimiter == "$" and "\n" in text[start_index:delimiter_index]:
            search_index = delimiter_index + len(delimiter)
            continue
        return delimiter_index
    return -1


def _split_inline_math_segments(text: str) -> list[dict[str, object]]:
    value = _normalize_line_endings(text)
    if not value:
        return []
    if "$" not in value:
        return [{"type": "text", "text": value}]

    segments: list[dict[str, object]] = []
    buffer: list[str] = []
    index = 0

    def flush_buffer() -> None:
        if buffer:
            segments.append({"type": "text", "text": "".join(buffer)})
            buffer.clear()

    while index < len(value):
        char = value[index]
        if char == "\\":
            next_char = value[index + 1] if index + 1 < len(value) else ""
            if next_char == "$":
                buffer.append("$")
                index += 2
                continue
            buffer.append(char)
            index += 1
            continue

        if char != "$":
            buffer.append(char)
            index += 1
            continue

        display_mode = index + 1 < len(value) and value[index + 1] == "$"
        delimiter = "$$" if display_mode else "$"
        math_start = index + len(delimiter)
        math_end = _find_math_delimiter(value, math_start, delimiter)
        if math_end < 0:
            buffer.append(char)
            index += 1
            continue

        math_text = value[math_start:math_end].strip()
        if not math_text:
            buffer.append(delimiter)
            index = math_start
            continue

        flush_buffer()
        segments.append({"type": "math", "text": math_text, "display": display_mode})
        index = math_end + len(delimiter)

    flush_buffer()
    return segments


_PDF_TABLE_TOTAL_WIDTH = 450.0  # pts, approx A4 content area with standard margins
_PDF_INLINE_BOLD_RE = re.compile(r"\*\*(.+?)\*\*|__(.+?)__", re.DOTALL)
_PDF_INLINE_ITALIC_RE = re.compile(
    r"(?<!\w)\*(?!\s)(.+?)(?<!\s)\*(?!\w)|(?<!\w)_(?!\s)(.+?)(?<!\s)_(?!\w)", re.DOTALL
)
_PDF_INLINE_CODE_RE = re.compile(r"`([^`]+)`")
_PDF_MATH_FRAC_RE = re.compile(r"\\frac\s*\{([^{}]+)\}\s*\{([^{}]+)\}")
_PDF_MATH_SQRT_RE = re.compile(r"\\sqrt\s*\{([^{}]+)\}")
_PDF_MATH_TEXT_RE = re.compile(r"\\text\s*\{([^{}]+)\}")

_PDF_MATH_REPLACEMENTS = {
    r"\\cdot": "·",
    r"\\times": "×",
    r"\\pm": "±",
    r"\\mp": "∓",
    r"\\neq": "≠",
    r"\\leq": "≤",
    r"\\geq": "≥",
    r"\\approx": "≈",
    r"\\sim": "∼",
    r"\\infty": "∞",
    r"\\rightarrow": "→",
    r"\\leftarrow": "←",
    r"\\leftrightarrow": "↔",
    r"\\Rightarrow": "⇒",
    r"\\Leftarrow": "⇐",
    r"\\sum": "∑",
    r"\\prod": "∏",
    r"\\int": "∫",
    r"\\partial": "∂",
    r"\\nabla": "∇",
    r"\\alpha": "α",
    r"\\beta": "β",
    r"\\gamma": "γ",
    r"\\delta": "δ",
    r"\\epsilon": "ε",
    r"\\theta": "θ",
    r"\\lambda": "λ",
    r"\\mu": "μ",
    r"\\pi": "π",
    r"\\sigma": "σ",
    r"\\phi": "φ",
    r"\\omega": "ω",
}

_PDF_SUPERSCRIPT_MAP = str.maketrans({
    "0": "⁰", "1": "¹", "2": "²", "3": "³", "4": "⁴", "5": "⁵", "6": "⁶", "7": "⁷", "8": "⁸", "9": "⁹",
    "+": "⁺", "-": "⁻", "=": "⁼", "(": "⁽", ")": "⁾", "n": "ⁿ", "i": "ⁱ",
})
_PDF_SUBSCRIPT_MAP = str.maketrans({
    "0": "₀", "1": "₁", "2": "₂", "3": "₃", "4": "₄", "5": "₅", "6": "₆", "7": "₇", "8": "₈", "9": "₉",
    "+": "₊", "-": "₋", "=": "₌", "(": "₍", ")": "₎", "a": "ₐ", "e": "ₑ", "h": "ₕ", "i": "ᵢ", "j": "ⱼ",
    "k": "ₖ", "l": "ₗ", "m": "ₘ", "n": "ₙ", "o": "ₒ", "p": "ₚ", "r": "ᵣ", "s": "ₛ", "t": "ₜ", "u": "ᵤ",
    "v": "ᵥ", "x": "ₓ",
})


def _format_pdf_math_script(script_kind: str, raw_value: str) -> str:
    value = str(raw_value or "").strip()
    if not value:
        return ""
    if script_kind == "sup" and re.fullmatch(r"[0-9+\-=()ni]+", value):
        return value.translate(_PDF_SUPERSCRIPT_MAP)
    if script_kind == "sub" and re.fullmatch(r"[0-9x]+", value):
        return value.translate(_PDF_SUBSCRIPT_MAP)
    prefix = "^" if script_kind == "sup" else "_"
    if len(value) == 1:
        return f"{prefix}{value}"
    return f"{prefix}({value})"


def _apply_pdf_inline_formatting(text: str, *, mono_font_name: str = "Courier") -> str:
    """Convert inline markdown formatting to ReportLab paragraph XML markup."""
    # Escape HTML entities first so bold/italic regexes only match raw markers
    value = html_escape(str(text or ""), quote=False)
    value = _PDF_INLINE_BOLD_RE.sub(lambda m: f"<b>{m.group(1) or m.group(2)}</b>", value)
    value = _PDF_INLINE_ITALIC_RE.sub(lambda m: f"<i>{m.group(1) or m.group(2)}</i>", value)
    value = _PDF_INLINE_CODE_RE.sub(lambda m: f'<font face="{mono_font_name}">{m.group(1)}</font>', value)
    return value


def _normalize_pdf_math_text(text: str) -> str:
    value = _normalize_line_endings(str(text or "").strip())
    if not value:
        return ""

    while True:
        next_value = _PDF_MATH_FRAC_RE.sub(lambda m: f"({m.group(1)})/({m.group(2)})", value)
        if next_value == value:
            break
        value = next_value

    while True:
        next_value = _PDF_MATH_SQRT_RE.sub(lambda m: f"√({m.group(1)})", value)
        if next_value == value:
            break
        value = next_value

    value = _PDF_MATH_TEXT_RE.sub(lambda m: m.group(1), value)
    value = value.replace(r"\left", "").replace(r"\right", "")
    for source, target in _PDF_MATH_REPLACEMENTS.items():
        value = value.replace(source, target)

    value = re.sub(r"\s+", " ", value).strip()
    return value


def _render_pdf_inline_markup(text: str, *, mono_font_name: str = "Courier") -> str:
    parts: list[str] = []
    for segment in _split_inline_math_segments(text):
        if str(segment.get("type") or "") == "math":
            math_text = _escape_pdf_text(_normalize_pdf_math_text(str(segment.get("text") or "")))
            parts.append(f'<font face="{mono_font_name}">{math_text}</font>')
        else:
            parts.append(_apply_pdf_inline_formatting(str(segment.get("text") or ""), mono_font_name=mono_font_name))
    return "".join(parts)


def _append_docx_inline_runs(paragraph, text: str) -> None:
    segments = _split_inline_math_segments(text)
    if not segments:
        paragraph.add_run(" ")
        return

    for segment in segments:
        run = paragraph.add_run(str(segment.get("text") or ""))
        if str(segment.get("type") or "") == "math":
            run.italic = True
            run.font.name = "Courier New"


def _clean_markdown_inline(text: str, *, preserve_formatting: bool = False) -> str:
    value = _normalize_line_endings(text)

    def _replace_image(match: re.Match) -> str:
        alt_text = (match.group(1) or "").strip()
        url = (match.group(2) or "").strip()
        return alt_text or url

    def _replace_link(match: re.Match) -> str:
        label = _clean_markdown_inline(match.group(1) or "", preserve_formatting=preserve_formatting)
        url = (match.group(2) or "").strip()
        if not label:
            return url
        if label == url:
            return label
        return f"{label} ({url})" if url else label

    value = re.sub(r"!\[([^\]]*)\]\(([^)]+)\)", _replace_image, value)
    value = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", _replace_link, value)
    if not preserve_formatting:
        value = re.sub(r"`([^`]+)`", lambda match: match.group(1), value)
        value = re.sub(r"(\*\*|__)(.+?)\1", lambda match: match.group(2), value)
        value = re.sub(r"(?<!\w)(\*|_)(?!\s)(.+?)(?<!\s)\1(?!\w)", lambda match: match.group(2), value)
    value = re.sub(r"~~(.+?)~~", lambda match: match.group(1), value)
    value = re.sub(r"<(https?://[^>]+)>", lambda match: match.group(1), value)
    value = value.replace("\\*", "*").replace("\\_", "_").replace("\\`", "`")
    value = value.replace("\\[", "[").replace("\\]", "]").replace("\\(", "(").replace("\\)", ")")
    value = re.sub(r"\s+", " ", value)
    return value.strip()


def _extract_list_item_texts(raw_items) -> list[str]:
    items: list[str] = []
    for item in raw_items or []:
        if isinstance(item, dict):
            item_text = str(item.get("text") or "").strip()
        else:
            item_text = str(item).strip()
        if item_text:
            items.append(item_text)
    return items


def _extract_ordered_list_start(block: dict[str, object]) -> int:
    try:
        start_value = int(block.get("start") or 1)
    except (TypeError, ValueError):
        return 1
    return max(1, start_value)


def _iter_markdown_blocks(text: str, *, preserve_inline_formatting: bool = False) -> list[dict[str, object]]:
    normalized = _normalize_line_endings(text)
    if not normalized.strip():
        return []

    blocks: list[dict[str, object]] = []
    paragraph_lines: list[str] = []
    list_items: list[dict[str, object]] = []
    list_kind: str | None = None
    list_start: int | None = None
    pending_list_break = False
    table_rows: list[list[str]] = []
    code_lines: list[str] = []
    in_code = False
    code_language = ""

    def flush_paragraph() -> None:
        nonlocal paragraph_lines
        if not paragraph_lines:
            return
        paragraph_text = _clean_markdown_inline(
            " ".join(part.strip() for part in paragraph_lines if part.strip()),
            preserve_formatting=preserve_inline_formatting,
        )
        paragraph_lines = []
        if paragraph_text:
            segments = _split_inline_math_segments(paragraph_text)
            if (
                len(segments) == 1
                and str(segments[0].get("type") or "") == "math"
                and bool(segments[0].get("display"))
            ):
                blocks.append({"type": "display_math", "text": str(segments[0].get("text") or "")})
            else:
                blocks.append({"type": "paragraph", "text": paragraph_text})

    def flush_list() -> None:
        nonlocal list_items, list_kind, list_start, pending_list_break
        if not list_items:
            return
        list_block: dict[str, object] = {
            "type": "list",
            "kind": list_kind or "bullet",
            "items": list_items[:],
        }
        if (list_kind or "bullet") == "ordered" and isinstance(list_start, int) and list_start > 0:
            list_block["start"] = list_start
        blocks.append(list_block)
        list_items = []
        list_kind = None
        list_start = None
        pending_list_break = False

    def resolve_pending_list_break(line: str) -> None:
        nonlocal pending_list_break
        if not pending_list_break:
            return
        continues_same_list = (list_kind == "bullet" and _MARKDOWN_BULLET_RE.match(line)) or (
            list_kind == "ordered" and _MARKDOWN_ORDERED_RE.match(line)
        )
        if not continues_same_list:
            flush_list()
        pending_list_break = False

    def flush_table() -> None:
        nonlocal table_rows
        if not table_rows:
            return
        if preserve_inline_formatting:
            cleaned: list[list[str]] = [
                [_clean_markdown_inline(cell or "", preserve_formatting=True) for cell in row]
                for row in table_rows
            ]
            blocks.append({"type": "table", "rows": cleaned})
        else:
            for row in table_rows:
                row_text = _clean_markdown_inline(" | ".join(cell for cell in row if cell is not None))
                if row_text:
                    blocks.append({"type": "paragraph", "text": row_text})
        table_rows = []

    for line in normalized.split("\n"):
        fence_match = _MARKDOWN_FENCE_RE.match(line)
        if fence_match:
            resolve_pending_list_break(line)
            if in_code:
                block_text = "\n".join(code_lines)
                if code_language in {"markdown", "md", "mkd"}:
                    flush_paragraph()
                    flush_list()
                    flush_table()
                    blocks.extend(_iter_markdown_blocks(block_text, preserve_inline_formatting=preserve_inline_formatting))
                else:
                    blocks.append({"type": "code", "text": block_text})
                code_lines = []
                code_language = ""
                in_code = False
            else:
                flush_paragraph()
                flush_list()
                flush_table()
                in_code = True
                code_language = (fence_match.group(1) or "").strip().lower()
            continue

        if in_code:
            code_lines.append(line)
            continue

        stripped = line.strip()
        if not stripped:
            flush_paragraph()
            flush_table()
            if pending_list_break:
                flush_list()
            elif list_items:
                pending_list_break = True
            else:
                flush_list()
            continue

        resolve_pending_list_break(line)

        if _MARKDOWN_HORIZONTAL_RULE_RE.match(line):
            flush_paragraph()
            flush_list()
            flush_table()
            blocks.append({"type": "spacer"})
            continue

        heading_match = _MARKDOWN_HEADING_RE.match(line)
        if heading_match:
            flush_paragraph()
            flush_list()
            flush_table()
            blocks.append(
                {
                    "type": "heading",
                    "level": len(heading_match.group(1)),
                    "text": _clean_markdown_inline(heading_match.group(2), preserve_formatting=preserve_inline_formatting),
                }
            )
            continue

        if _MARKDOWN_TABLE_SEPARATOR_RE.match(line):
            continue

        if _MARKDOWN_TABLE_ROW_RE.match(line) and line.strip().count("|") >= 2:
            flush_paragraph()
            flush_list()
            cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
            table_rows.append(cells)
            continue

        bullet_match = _MARKDOWN_BULLET_RE.match(line)
        if bullet_match:
            flush_paragraph()
            flush_table()
            if list_kind not in {None, "bullet"}:
                flush_list()
            list_kind = "bullet"
            list_items.append(
                {
                    "text": _clean_markdown_inline(
                        bullet_match.group(1),
                        preserve_formatting=preserve_inline_formatting,
                    )
                }
            )
            continue

        ordered_match = _MARKDOWN_ORDERED_RE.match(line)
        if ordered_match:
            flush_paragraph()
            flush_table()
            if list_kind not in {None, "ordered"}:
                flush_list()
            list_kind = "ordered"
            explicit_number = int(ordered_match.group(1)) if ordered_match.group(1) else 1
            if list_start is None:
                list_start = max(1, explicit_number)
            list_items.append(
                {
                    "text": _clean_markdown_inline(
                        ordered_match.group(2),
                        preserve_formatting=preserve_inline_formatting,
                    ),
                    "number": max(1, explicit_number),
                }
            )
            continue

        quote_match = _MARKDOWN_QUOTE_RE.match(line)
        if quote_match:
            flush_paragraph()
            flush_list()
            flush_table()
            quote_text = _clean_markdown_inline(quote_match.group(1), preserve_formatting=preserve_inline_formatting)
            if quote_text:
                blocks.append({"type": "quote", "text": quote_text})
            continue

        flush_table()
        paragraph_lines.append(line)

    if in_code:
        block_text = "\n".join(code_lines)
        if code_language in {"markdown", "md", "mkd"}:
            blocks.extend(_iter_markdown_blocks(block_text, preserve_inline_formatting=preserve_inline_formatting))
        else:
            blocks.append({"type": "code", "text": block_text})

    flush_paragraph()
    if pending_list_break:
        flush_list()
    flush_list()
    flush_table()
    return blocks


def append_markdown_pdf_story(
    story: list,
    markdown_text: str,
    *,
    body_style,
    heading1_style,
    heading_style,
    subheading_style,
    code_style,
    quote_style=None,
    math_style=None,
    table_header_style=None,
    table_body_style=None,
    mono_font_name: str = "Courier",
    heading_level_offset: int = 0,
    empty_text: str = "(empty)",
) -> None:
    blocks = _iter_markdown_blocks(markdown_text, preserve_inline_formatting=True)
    if not blocks:
        story.append(Paragraph(_escape_pdf_text(empty_text), body_style))
        return

    quote_style = quote_style or body_style
    math_style = math_style or body_style
    table_header_style = table_header_style or body_style
    table_body_style = table_body_style or body_style

    offset = max(0, int(heading_level_offset or 0))
    for block in blocks:
        block_type = str(block.get("type") or "")
        if block_type == "heading":
            level = min(6, max(1, int(block.get("level") or 1) + offset))
            style = heading1_style if level <= 1 else heading_style if level == 2 else subheading_style
            story.append(Paragraph(_render_pdf_inline_markup(str(block.get("text") or "Untitled"), mono_font_name=mono_font_name), style))
        elif block_type == "paragraph":
            paragraph_text = str(block.get("text") or "").strip()
            if paragraph_text:
                story.append(Paragraph(_render_pdf_inline_markup(paragraph_text, mono_font_name=mono_font_name), body_style))
        elif block_type == "quote":
            quote_text = str(block.get("text") or "").strip()
            if quote_text:
                story.append(Paragraph(_render_pdf_inline_markup(quote_text, mono_font_name=mono_font_name), quote_style))
        elif block_type == "display_math":
            math_text = _normalize_pdf_math_text(str(block.get("text") or ""))
            if math_text:
                story.append(Paragraph(_escape_pdf_text(math_text), math_style))
        elif block_type == "list":
            items = _extract_list_item_texts(block.get("items") or [])
            if items:
                is_ordered = str(block.get("kind") or "bullet") == "ordered"
                if is_ordered:
                    ordered_style = ParagraphStyle(
                        f"{getattr(body_style, 'name', 'Body')}OrderedList",
                        parent=body_style,
                        leftIndent=max(18, float(getattr(body_style, "leftIndent", 0)) + 18),
                        firstLineIndent=-14,
                        spaceAfter=3,
                    )
                    start_value = _extract_ordered_list_start(block)
                    for index, item in enumerate(items):
                        prefix = f"{start_value + index}. "
                        story.append(
                            Paragraph(
                                _render_pdf_inline_markup(f"{prefix}{item}", mono_font_name=mono_font_name),
                                ordered_style,
                            )
                        )
                    story.append(Spacer(1, 4))
                else:
                    story.append(
                        ListFlowable(
                            [ListItem(Paragraph(_render_pdf_inline_markup(item, mono_font_name=mono_font_name), body_style)) for item in items],
                            bulletType="bullet",
                            leftIndent=20,
                            bulletFontName=getattr(body_style, "fontName", "Helvetica"),
                            bulletFontSize=max(8, float(getattr(body_style, "fontSize", 10))),
                            bulletColor=getattr(body_style, "textColor", colors.black),
                            bulletDedent=8,
                        )
                    )
                    story.append(Spacer(1, 6))
        elif block_type == "code":
            code_text = str(block.get("text") or "").rstrip()
            if code_text:
                story.append(Preformatted(code_text, code_style))
                story.append(Spacer(1, 4))
        elif block_type == "table":
            table_rows_data = block.get("rows") or []
            if table_rows_data:
                num_cols = max((len(r) for r in table_rows_data), default=1)
                col_width = _PDF_TABLE_TOTAL_WIDTH / num_cols
                table_data = []
                for row_index, row in enumerate(table_rows_data):
                    cell_style = table_header_style if row_index == 0 and len(table_rows_data) > 1 else table_body_style
                    table_data.append(
                        [
                            Paragraph(_render_pdf_inline_markup(str(cell or ""), mono_font_name=mono_font_name), cell_style)
                            for cell in row
                        ]
                    )
                rendered_table = Table(
                    table_data, colWidths=[col_width] * num_cols, hAlign="LEFT", repeatRows=1 if len(table_rows_data) > 1 else 0
                )
                rendered_table.setStyle(
                    TableStyle([
                        ("BOX", (0, 0), (-1, -1), 0.7, colors.HexColor("#d7deea")),
                        ("INNERGRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#dde4ef")),
                        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#edf3ff")),
                        ("TOPPADDING", (0, 0), (-1, -1), 6),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                        ("LEFTPADDING", (0, 0), (-1, -1), 8),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f9fbfe")]),
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ])
                )
                story.append(rendered_table)
                story.append(Spacer(1, 8))
        elif block_type == "spacer":
            story.append(Spacer(1, 4))


def append_markdown_docx(
    document: Document,
    markdown_text: str,
    *,
    heading_level_offset: int = 0,
    empty_text: str = "(empty)",
) -> None:
    blocks = _iter_markdown_blocks(markdown_text)
    if not blocks:
        document.add_paragraph(empty_text)
        return

    offset = max(0, int(heading_level_offset or 0))
    for block in blocks:
        block_type = str(block.get("type") or "")
        if block_type == "heading":
            level = min(9, max(1, int(block.get("level") or 1) + offset))
            paragraph = document.add_paragraph(style=f"Heading {level}")
            _append_docx_inline_runs(paragraph, str(block.get("text") or "Untitled"))
        elif block_type == "paragraph":
            paragraph_text = str(block.get("text") or "").strip()
            paragraph = document.add_paragraph()
            _append_docx_inline_runs(paragraph, paragraph_text or " ")
        elif block_type == "list":
            items = _extract_list_item_texts(block.get("items") or [])
            is_ordered = str(block.get("kind") or "bullet") == "ordered"
            start_value = _extract_ordered_list_start(block)
            for index, item_text in enumerate(items):
                if not item_text:
                    continue
                try:
                    if is_ordered:
                        paragraph = document.add_paragraph(style="List Paragraph")
                        _append_docx_inline_runs(paragraph, f"{start_value + index}. {item_text}")
                    else:
                        paragraph = document.add_paragraph(style="List Bullet")
                        _append_docx_inline_runs(paragraph, item_text)
                except Exception:
                    prefix = f"{start_value + index}. " if is_ordered else "- "
                    paragraph = document.add_paragraph()
                    _append_docx_inline_runs(paragraph, f"{prefix}{item_text}")
        elif block_type == "code":
            code_text = str(block.get("text") or "").rstrip()
            if code_text:
                paragraph = document.add_paragraph()
                run = paragraph.add_run(code_text)
                run.font.name = "Courier New"
        elif block_type == "spacer":
            document.add_paragraph(" ")