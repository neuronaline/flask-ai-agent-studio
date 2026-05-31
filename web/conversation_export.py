from __future__ import annotations

from datetime import datetime, timezone
from io import BytesIO
import json

from utils.export_styles import MONO_FONT, build_print_pdf_styles
from lib.markdown_rendering import append_markdown_docx, append_markdown_pdf_story

from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

from services.canvas import extract_canvas_documents
from core.db import extract_message_attachments


def _escape_pdf_text(value: str) -> str:
    return str(value or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _safe_text(value, fallback: str = "") -> str:
    text = str(value or "").strip()
    return text or fallback


def _append_detail_line(lines: list[str], label: str, value) -> None:
    text = _safe_text(value)
    if text:
        lines.append(f"- {label}: {text}")


def _append_boolean_detail_line(lines: list[str], label: str, value) -> None:
    if value is True:
        lines.append(f"- {label}: Yes")
    elif value is False:
        lines.append(f"- {label}: No")


def _join_values(values: list[str]) -> str:
    return ", ".join(value for value in values if value)


def _build_export_header(conversation: dict, message_count: int) -> list[str]:
    title = str(conversation.get("title") or "Conversation Export").strip() or "Conversation Export"
    model = str(conversation.get("model") or "").strip()
    conversation_id = conversation.get("id")
    exported_at = datetime.now(timezone.utc).astimezone().isoformat(timespec="seconds")
    lines = [f"# {title}", "", f"Exported at: {exported_at}"]
    if conversation_id is not None:
        lines.append(f"Conversation ID: {conversation_id}")
    lines.append(f"Message count: {message_count}")
    if model:
        lines.append(f"Model: {model}")
    return lines


def _build_message_metadata_details(message: dict) -> str | None:
    lines = []
    _append_detail_line(lines, "Message ID", message.get("id"))
    _append_detail_line(lines, "Position", message.get("position"))
    _append_detail_line(lines, "Created at", message.get("created_at"))
    _append_detail_line(lines, "Tool call ID", message.get("tool_call_id"))
    _append_detail_line(lines, "Prompt tokens", message.get("prompt_tokens"))
    _append_detail_line(lines, "Completion tokens", message.get("completion_tokens"))
    _append_detail_line(lines, "Total tokens", message.get("total_tokens"))

    usage = message.get("usage") if isinstance(message.get("usage"), dict) else None
    if usage:
        usage_lines = []
        _append_detail_line(usage_lines, "Prompt tokens", usage.get("prompt_tokens"))
        _append_detail_line(usage_lines, "Completion tokens", usage.get("completion_tokens"))
        _append_detail_line(usage_lines, "Total tokens", usage.get("total_tokens"))
        _append_detail_line(usage_lines, "Estimated input tokens", usage.get("estimated_input_tokens"))
        if usage_lines:
            lines.extend(["", "Usage"])
            lines.extend(usage_lines)

    return "\n".join(lines).strip() or None


def _build_canvas_documents_details(canvas_documents: list[dict]) -> str | None:
    lines = []
    for document in canvas_documents:
        if not isinstance(document, dict):
            continue
        title = _safe_text(document.get("title"), fallback="Canvas")
        lines.append(f"### Canvas: {title}")
        doc_content = _safe_text(document.get("content"))
        if doc_content:
            format_name = str(document.get("format") or "").strip().lower()
            fence_name = _safe_text(document.get("language"), fallback="text") if format_name == "code" else "markdown"
            lines.extend(["", f"```{fence_name}", doc_content, "```", ""])
    return "\n".join(lines).strip() or None


def _build_attachment_details(attachments: list[dict]) -> str | None:
    blocks = []
    for index, attachment in enumerate(attachments, start=1):
        if not isinstance(attachment, dict):
            continue

        kind = _safe_text(attachment.get("kind"), fallback="attachment")
        title = _safe_text(
            attachment.get("file_name")
            or attachment.get("image_name")
            or attachment.get("video_title")
            or attachment.get("file_id")
            or attachment.get("image_id")
            or attachment.get("video_id"),
            fallback="Attachment",
        )
        heading = f"#### Attachment {index}: {kind.title()}"
        if title and title.lower() != kind.lower():
            heading = f"{heading} - {title}"

        lines = [heading]
        if kind == "image":
            _append_detail_line(lines, "Image ID", attachment.get("image_id"))
            _append_detail_line(lines, "Name", attachment.get("image_name"))
            _append_detail_line(lines, "MIME type", attachment.get("image_mime_type"))
            _append_detail_line(lines, "Analysis method", attachment.get("analysis_method"))
            _append_detail_line(lines, "OCR text", attachment.get("ocr_text"))
            _append_detail_line(lines, "Vision summary", attachment.get("vision_summary"))
            _append_detail_line(lines, "Assistant guidance", attachment.get("assistant_guidance"))
            key_points = attachment.get("key_points") if isinstance(attachment.get("key_points"), list) else []
            if key_points:
                lines.append("- Key points:")
                for point in key_points[:8]:
                    point_text = _safe_text(point)
                    if point_text:
                        lines.append(f"  - {point_text}")
        elif kind == "video":
            _append_detail_line(lines, "Video ID", attachment.get("video_id"))
            _append_detail_line(lines, "Title", attachment.get("video_title"))
            _append_detail_line(lines, "URL", attachment.get("video_url"))
            _append_detail_line(lines, "Platform", attachment.get("video_platform"))
            _append_detail_line(lines, "Transcript language", attachment.get("transcript_language"))
            _append_detail_line(lines, "Transcript context", attachment.get("transcript_context_block"))
            _append_boolean_detail_line(lines, "Transcript truncated", attachment.get("transcript_text_truncated"))
        else:
            _append_detail_line(lines, "File ID", attachment.get("file_id"))
            _append_detail_line(lines, "File name", attachment.get("file_name"))
            _append_detail_line(lines, "MIME type", attachment.get("file_mime_type"))
            _append_detail_line(lines, "Submission mode", attachment.get("submission_mode"))
            _append_detail_line(lines, "Canvas mode", attachment.get("canvas_mode"))
            _append_detail_line(lines, "File context", attachment.get("file_context_block"))
            _append_boolean_detail_line(lines, "File truncated", attachment.get("file_text_truncated"))
            page_ids = attachment.get("visual_page_image_ids") if isinstance(attachment.get("visual_page_image_ids"), list) else []
            if page_ids:
                _append_detail_line(lines, "Visual page image IDs", _join_values([_safe_text(page_id) for page_id in page_ids[:8]]))
            _append_detail_line(lines, "Visual page count", attachment.get("visual_page_count"))
            _append_detail_line(lines, "Visual total page count", attachment.get("visual_total_page_count"))
            _append_detail_line(lines, "Visual page limit", attachment.get("visual_page_limit"))
            _append_boolean_detail_line(lines, "Visual pages truncated", attachment.get("visual_pages_truncated"))

        blocks.append("\n".join(lines).strip())
    return "\n\n".join(blocks).strip() or None


def _iter_message_sections(messages: list[dict]) -> list[dict]:
    sections = []
    section_index = 0
    for message in messages or []:
        if not isinstance(message, dict):
            continue

        role = str(message.get("role") or "message").strip() or "message"
        content = str(message.get("content") or "").strip()
        metadata = message.get("metadata") if isinstance(message.get("metadata"), dict) else {}
        details = []

        message_metadata = _build_message_metadata_details(message)
        if message_metadata:
            details.append(("Message Metadata", message_metadata))

        reasoning_content = str(metadata.get("reasoning_content") or "").strip()
        if role == "assistant" and reasoning_content:
            details.append(("Reasoning", reasoning_content))

        tool_trace = metadata.get("tool_trace") if isinstance(metadata.get("tool_trace"), list) else []
        if tool_trace:
            lines = []
            for entry in tool_trace:
                if not isinstance(entry, dict):
                    continue
                tool_name = str(entry.get("tool_name") or "tool").strip() or "tool"
                step = entry.get("step")
                summary = str(entry.get("summary") or entry.get("preview") or "").strip()
                prefix = f"Step {int(step)}" if isinstance(step, (int, float)) else "Step"
                lines.append(f"- {prefix}: {tool_name}{(': ' + summary) if summary else ''}")
            if lines:
                details.append(("Tool Trace", "\n".join(lines)))

        tool_results = metadata.get("tool_results") if isinstance(metadata.get("tool_results"), list) else []
        if tool_results:
            lines = []
            for entry in tool_results:
                if not isinstance(entry, dict):
                    continue
                tool_name = str(entry.get("tool_name") or "tool").strip() or "tool"
                summary = str(entry.get("summary") or "").strip()
                lines.append(f"- {tool_name}{(': ' + summary) if summary else ''}")
            if lines:
                details.append(("Tool Results", "\n".join(lines)))

        attachments = extract_message_attachments(metadata)
        if attachments:
            attachment_details = _build_attachment_details(attachments)
            if attachment_details:
                details.append(("Attachments", attachment_details))

        canvas_documents = extract_canvas_documents(metadata)
        if canvas_documents:
            canvas_details = _build_canvas_documents_details(canvas_documents)
            if canvas_details:
                details.append(("Canvas Documents", canvas_details))

        has_non_reasoning_details = any(label not in {"Message Metadata", "Reasoning"} for label, _ in details)
        if role == "assistant" and not content and not has_non_reasoning_details:
            continue

        section_index += 1

        sections.append(
            {
                "title": f"## {section_index}. {role.title()}",
                "content": content,
                "details": details,
            }
        )
    return sections


def _build_raw_export_message_metadata(message: dict) -> dict | None:
    metadata = message.get("metadata") if isinstance(message.get("metadata"), dict) else {}
    if not metadata:
        return None

    cleaned = {
        key: value
        for key, value in metadata.items()
        if key != "_edit_replay_deleted"
    }
    return cleaned or None


def _build_raw_export_transcript(messages: list[dict]) -> list[dict]:
    transcript = []
    for message in messages or []:
        if not isinstance(message, dict):
            continue
        metadata = message.get("metadata") if isinstance(message.get("metadata"), dict) else {}
        raw_metadata = _build_raw_export_message_metadata(message)
        transcript.append(
            {
                "id": message.get("id"),
                "position": message.get("position"),
                "role": _safe_text(message.get("role")),
                "content": str(message.get("content") or ""),
                "tool_call_id": _safe_text(message.get("tool_call_id")) or None,
                "tool_calls": message.get("tool_calls") if isinstance(message.get("tool_calls"), list) else [],
                "has_reasoning": bool(str(metadata.get("reasoning_content") or "").strip()),
                "usage": message.get("usage") if isinstance(message.get("usage"), dict) else None,
                "created_at": _safe_text(message.get("created_at")) or None,
                "deleted_at": _safe_text(message.get("deleted_at")) or None,
                "metadata": raw_metadata,
            }
        )
    return transcript


def _build_raw_export_capture_status(messages: list[dict], invocations: list[dict]) -> dict:
    if invocations:
        status = "available"
    elif any(str((message or {}).get("role") or "").strip() == "assistant" for message in messages or []):
        status = "unavailable_for_legacy_conversation"
    else:
        status = "no_completed_assistant_turns"
    return {
        "status": status,
        "invocation_count": len(invocations or []),
        "future_only_exact_snapshots": True,
    }


def build_conversation_json_download(conversation: dict, messages: list[dict], invocations: list[dict]) -> bytes:
    payload = {
        "export_type": "conversation_raw_model_invocations",
        "schema_version": 1,
        "exported_at": datetime.now(timezone.utc).astimezone().isoformat(timespec="seconds"),
        "conversation": {
            "id": conversation.get("id"),
            "title": _safe_text(conversation.get("title"), fallback="Conversation Export"),
            "model": _safe_text(conversation.get("model")) or None,
            "created_at": _safe_text(conversation.get("created_at")) or None,
            "updated_at": _safe_text(conversation.get("updated_at")) or None,
            "message_count": len(messages or []),
        },
        "capture_status": _build_raw_export_capture_status(messages, invocations),
        "capture_scope": [
            "main agent model calls",
            "sub-agent model calls",
            "fetch_url_summarized helper-model calls",
        ],
        "limitations": [
            "Exact snapshots are only available for turns captured after raw invocation logging was enabled.",
            "Legacy conversations may contain transcript history without exact provider request snapshots.",
        ],
        "transcript": _build_raw_export_transcript(messages),
        "invocations": invocations or [],
    }
    return json.dumps(payload, ensure_ascii=False, indent=2).encode("utf-8")


def build_conversation_markdown_download(conversation: dict, messages: list[dict]) -> bytes:
    sections = _iter_message_sections(messages)
    lines = _build_export_header(conversation, len(sections))
    for section in sections:
        lines.extend(["", section["title"], ""])
        lines.append(section["content"] or "_(empty)_")
        for label, value in section["details"]:
            lines.extend(["", f"### {label}", "", value])
    lines.append("")
    return "\n".join(lines).encode("utf-8")


def build_conversation_docx_download(conversation: dict, messages: list[dict]) -> bytes:
    sections = _iter_message_sections(messages)
    document = Document()
    title = str(conversation.get("title") or "Conversation Export").strip() or "Conversation Export"
    document.add_heading(title, level=0)
    exported_at = datetime.now(timezone.utc).astimezone().isoformat(timespec="seconds")
    meta_parts = [f"Exported at: {exported_at}", f"Message count: {len(sections)}"]
    conversation_id = conversation.get("id")
    if conversation_id is not None:
        meta_parts.insert(1, f"Conversation ID: {conversation_id}")
    model = str(conversation.get("model") or "").strip()
    if model:
        meta_parts.append(f"Model: {model}")
    document.add_paragraph(" | ".join(meta_parts))

    for section in sections:
        document.add_heading(section["title"].replace("## ", ""), level=1)
        append_markdown_docx(document, section["content"] or "(empty)", heading_level_offset=2)
        for label, value in section["details"]:
            document.add_heading(label, level=2)
            append_markdown_docx(document, str(value or "").strip() or "(empty)", heading_level_offset=2)

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def build_conversation_pdf_download(conversation: dict, messages: list[dict]) -> bytes:
    sections = _iter_message_sections(messages)
    pdf_styles = build_print_pdf_styles("ConversationExport")
    title_style = pdf_styles["title"]
    heading_style = pdf_styles["heading"]
    subheading_style = pdf_styles["subheading"]
    body_style = pdf_styles["body"]
    meta_style = pdf_styles["meta"]
    code_style = pdf_styles["code"]
    quote_style = pdf_styles["quote"]
    math_style = pdf_styles["math"]
    table_header_style = pdf_styles["table_header"]
    table_body_style = pdf_styles["table_body"]

    title = str(conversation.get("title") or "Conversation Export").strip() or "Conversation Export"
    exported_at = datetime.now(timezone.utc).astimezone().isoformat(timespec="seconds")
    story = [Paragraph(_escape_pdf_text(title), title_style), Spacer(1, 6)]
    meta_parts = [f"Exported at: {exported_at}", f"Message count: {len(sections)}"]
    conversation_id = conversation.get("id")
    if conversation_id is not None:
        meta_parts.insert(1, f"Conversation ID: {conversation_id}")
    model = str(conversation.get("model") or "").strip()
    if model:
        meta_parts.append(f"Model: {model}")
    story.append(Paragraph(_escape_pdf_text(" | ".join(meta_parts)), meta_style))

    for section in sections:
        story.append(Spacer(1, 6))
        story.append(Paragraph(_escape_pdf_text(section["title"].replace("## ", "")), heading_style))
        append_markdown_pdf_story(
            story,
            section["content"] or "(empty)",
            body_style=body_style,
            heading1_style=subheading_style,
            heading_style=subheading_style,
            subheading_style=subheading_style,
            code_style=code_style,
            quote_style=quote_style,
            math_style=math_style,
            table_header_style=table_header_style,
            table_body_style=table_body_style,
            mono_font_name=MONO_FONT,
            heading_level_offset=2,
        )
        for label, value in section["details"]:
            story.append(Paragraph(_escape_pdf_text(label), subheading_style))
            append_markdown_pdf_story(
                story,
                str(value or "").strip() or "(empty)",
                body_style=body_style,
                heading1_style=subheading_style,
                heading_style=subheading_style,
                subheading_style=subheading_style,
                code_style=code_style,
                quote_style=quote_style,
                math_style=math_style,
                table_header_style=table_header_style,
                table_body_style=table_body_style,
                mono_font_name=MONO_FONT,
                heading_level_offset=2,
            )

    output = BytesIO()
    doc = SimpleDocTemplate(output, pagesize=A4, topMargin=18 * mm, bottomMargin=18 * mm)
    doc.build(story)
    return output.getvalue()